from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ---------- COVER PAGE ----------
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Community Grievance Redressal and\nPublic Service Tracking System')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NagarSeva — Civic Grievance Portal')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xE6, 0x8A, 0x00)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Course: UCS310 — Database Management Systems\nB.Tech, 2nd Year — Computer Engineering\nSub-Group: 2C16')
run.font.size = Pt(12)

doc.add_paragraph()

# Group members table
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Name'
table.cell(0, 1).text = 'Roll Number'
table.cell(1, 0).text = 'Vanshika'
table.cell(1, 1).text = '102403163'
table.cell(2, 0).text = 'Aishwarya Verma'
table.cell(2, 1).text = '102403170'
table.cell(3, 0).text = 'Diya Jain'
table.cell(3, 1).text = '102403154'
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted To: Dr. Simranjit Kaur')
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Thapar Institute of Engineering and Technology\n(Deemed to be University)\nPatiala')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('May 2026')
run.font.size = Pt(11)

doc.add_page_break()

# ---------- HELPER FUNCTIONS ----------
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    return p

def add_code_block(code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('[ INSERT SCREENSHOT OF SQL CODE HERE ]')
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p

def add_table_from_data(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ---------- READ THE MARKDOWN ----------
with open(r'c:\Users\diyaj\OneDrive\Desktop\dbms project\DBMS_Project_Report.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buffer = []
in_table = False
table_rows = []
table_headers = []
skip_cover = True  # skip the markdown cover page

while i < len(lines):
    line = lines[i].rstrip('\n').rstrip('\r')
    
    # Skip the markdown cover page (until first ## section)
    if skip_cover:
        if line.startswith('## 1.') or line.startswith('## Table of Contents'):
            if 'Table of Contents' in line:
                # Skip TOC entirely, we already have it implicitly
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('## 1.'):
                    i += 1
                continue
            skip_cover = False
        else:
            i += 1
            continue
    
    # Code blocks
    if line.strip().startswith('```'):
        if in_code:
            add_code_block('\n'.join(code_buffer))
            code_buffer = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue
    
    if in_code:
        code_buffer.append(line)
        i += 1
        continue
    
    # Tables
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= {'-', ':', ' '} for c in cells):
            # separator row, skip
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
            table_rows = []
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_data(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
    
    # Skip page break divs
    if 'page-break' in line or line.strip() == '<div' or line.strip() == '</div>':
        i += 1
        continue
    
    # Horizontal rules
    if line.strip() == '---':
        i += 1
        continue
    
    # Headings
    if line.startswith('## '):
        add_heading(line[3:].strip(), level=1)
        doc.add_paragraph()  # spacing
        i += 1
        continue
    if line.startswith('### '):
        add_heading(line[4:].strip(), level=2)
        i += 1
        continue
    if line.startswith('#### '):
        add_heading(line[5:].strip(), level=3)
        i += 1
        continue
    
    # Bullet points
    if line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        # Check for **bold prefix:**
        m = re.match(r'\*\*(.+?)\*\*\s*(.*)', text)
        if m:
            add_bullet(m.group(2), bold_prefix=m.group(1))
        else:
            # Strip remaining markdown bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_bullet(text)
        i += 1
        continue
    
    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.+)', line)
    if m:
        text = m.group(2).strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Number')
        p.add_run(text)
        i += 1
        continue
    
    # Empty lines
    if line.strip() == '':
        i += 1
        continue
    
    # Regular paragraphs
    text = line.strip()
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # italic
    text = re.sub(r'`(.+?)`', r'\1', text)  # inline code
    if text:
        add_body(text)
    
    i += 1

# Flush remaining table
if in_table:
    add_table_from_data(table_headers, table_rows)

# ---------- ADD HEADERS/FOOTERS ----------
for section in doc.sections:
    header = section.header
    hp = header.paragraphs[0]
    hp.text = 'Community Grievance Redressal and Public Service Tracking System — UCS310'
    hp.style = doc.styles['Header']
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = 'Vanshika · Aishwarya Verma · Diya Jain | TIET, Patiala'
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# Save
output_path = r'c:\Users\diyaj\OneDrive\Desktop\dbms project\DBMS_Project_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
